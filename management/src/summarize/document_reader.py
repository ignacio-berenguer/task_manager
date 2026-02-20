"""Document content extraction for various file formats."""

import logging
from pathlib import Path

logger = logging.getLogger('portfolio_summarize')


class UnsupportedFormatError(Exception):
    """Raised when a document format is not supported."""
    pass


def read_document(file_path: str) -> str:
    """
    Read document content based on file extension.

    Args:
        file_path: Path to the document file

    Returns:
        Extracted text content

    Raises:
        FileNotFoundError: If the file does not exist
        UnsupportedFormatError: If the file format is not supported
        Exception: For other read errors (permissions, encoding, etc.)
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")

    extension = path.suffix.lower()

    readers = {
        '.pdf': _read_pdf,
        '.docx': _read_docx,
        '.docm': _read_docx,
        '.doc': _read_doc,
        '.txt': _read_txt,
        '.pptx': _read_pptx,
    }

    reader = readers.get(extension)
    if not reader:
        raise UnsupportedFormatError(
            f"Unsupported file format '{extension}' for: {path.name}"
        )

    logger.debug(f"Reading {extension} file: {path.name}")
    return reader(path)


def _read_pdf(path: Path) -> str:
    """Extract text from a PDF file using pymupdf."""
    import pymupdf

    text_parts = []
    with pymupdf.open(str(path)) as doc:
        for page in doc:
            text_parts.append(page.get_text())

    text = "\n".join(text_parts).strip()
    if not text:
        raise ValueError(f"No text content extracted from PDF: {path.name}")
    return text


def _read_docx(path: Path) -> str:
    """Extract text from a DOCX file using python-docx, including tables."""
    from docx import Document

    doc = Document(str(path))

    # Extract paragraphs
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # Extract tables (cover page metadata, data tables, etc.)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        raise ValueError(f"No text content extracted from DOCX: {path.name}")
    return text


def _read_doc(path: Path) -> str:
    """Extract text from a legacy .doc file using pymupdf."""
    import pymupdf

    text_parts = []
    with pymupdf.open(str(path)) as doc:
        for page in doc:
            text_parts.append(page.get_text())

    text = "\n".join(text_parts).strip()
    if not text:
        raise ValueError(f"No text content extracted from DOC: {path.name}")
    return text


def _read_pptx(path: Path) -> str:
    """Extract text from a PPTX file using pymupdf."""
    import pymupdf

    text_parts = []
    with pymupdf.open(str(path)) as doc:
        for page in doc:
            text_parts.append(page.get_text())

    text = "\n".join(text_parts).strip()
    if not text:
        raise ValueError(f"No text content extracted from PPTX: {path.name}")
    return text


def _read_txt(path: Path) -> str:
    """Read a plain text file with encoding fallback."""
    for encoding in ('utf-8', 'latin-1'):
        try:
            text = path.read_text(encoding=encoding).strip()
            if text:
                return text
        except UnicodeDecodeError:
            continue

    raise ValueError(f"Could not read text file with any supported encoding: {path.name}")
